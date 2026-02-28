# export_to_word.py
# Generates a fully structured academic Word document (.docx)
# for the London COVID-19 Data Analysis project
# University of Essex — MSc Artificial Intelligence

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, colour=None):
    run.font.name        = name
    run.font.size        = Pt(size)
    run.font.bold        = bold
    run.font.italic      = italic
    if colour:
        run.font.color.rgb = RGBColor(*colour)

def heading(doc, text, level=1, font_size=14, colour=(0, 0, 0), space_before=12):
    p   = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(font_size)
        run.font.color.rgb = RGBColor(*colour)
        run.font.bold   = True
    return p

def body(doc, text, indent=False, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, italic=italic)
    return p

def code_block(doc, code_text):
    """Render code in a monospace, shaded style."""
    # Add a light grey shaded paragraph for each line
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.8)
        # shade background
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
    doc.add_paragraph()   # breathing space after block

def figure(doc, img_path, caption_text, fig_num, width_cm=15.5):
    """Embed an image with a numbered caption."""
    if not os.path.exists(img_path):
        body(doc, f"[Image not found: {img_path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"Figure {fig_num}: {caption_text}")
    set_font(r, size=10, italic=True)

def divider(doc):
    p = doc.add_paragraph("─" * 85)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.size  = Pt(7)
        run.font.color.rgb = RGBColor(180, 180, 180)

def read_source(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

# ── Source files & images ─────────────────────────────────────────────────────

SRC = {
    "acquire" : "Unit_5_covid_analysis.py",
    "clean"   : "Unit_5.2_clean_data.py",
    "analysis": "Unit_5.3_analysis.py",
    "simple"  : "Unit_5.3.1_simplified.py",
    "visuals" : "Unit_5.4_visualisations.py",
}

# Figures selected for the document (path, caption)
FIGURES = [
    ("london_covid_cleaning_overview.png",
     "Data Cleaning Overview — Missing Values, Case & Death Trends, and Case Fatality Rate"),
    ("chart_3_cases_trend.png",
     "New COVID-19 Cases Over Time with 4-Week Rolling Average, Mean and Median (London)"),
    ("chart_4_deaths_trend.png",
     "New COVID-19 Deaths (28-day basis) Over Time with 4-Week Rolling Average (London)"),
    ("chart_1_mean_vs_median.png",
     "Mean vs Median Comparison for Key COVID-19 Metrics — London, Mar 2020 – Apr 2022"),
    ("chart_2_mean_by_year.png",
     "Average Weekly Metrics Grouped by Year (2020, 2021, 2022) — London"),
    ("chart_5_hospital_admissions_trend.png",
     "Hospital Cases and New Admissions Trends Over Time — London (Dual-Axis Line Chart)"),
    ("chart_6_cfr_trend.png",
     "Case Fatality Rate (%) Over Time with Pre- and Post-Vaccination Eras Highlighted — London"),
    ("chart_7_monthly_cases_bar.png",
     "Monthly Total New Cases with Month-on-Month Percentage Change — London"),
    ("chart_8_yearly_totals.png",
     "Year-on-Year Total New Cases and Total Deaths — London"),
    ("chart_9_correlation_heatmap.png",
     "Pearson Correlation Heatmap Across Key COVID-19 Metrics — London"),
    ("chart_10_correlation_bars.png",
     "Pairwise Correlation Coefficients (r) Between COVID-19 Metrics — London"),
    ("chart_11_scatter_lag_correlation.png",
     "Scatter Plot: New Deaths vs New Cases with 2-Week Lag and Linear Trend Line — London"),
]

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

for txt, sz, bold in [
    ("University of Essex",               14, False),
    ("School of Computer Science and Electronic Engineering", 12, False),
    ("MSc Artificial Intelligence",       13, True),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    set_font(r, size=sz, bold=bold)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("COVID-19 Data Analysis Report")
set_font(r, size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("London Region — March 2020 to April 2022")
set_font(r, size=14, italic=True)

doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ("Module",   "Unit 5 — Data Analysis and Visualisation"),
    ("Dataset",  "UK Government COVID-19 Dashboard (London Region)"),
    ("Date",     datetime.date.today().strftime("%d %B %Y")),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(value)
    set_font(r2, size=12)

doc.add_page_break()

# ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────────

heading(doc, "Table of Contents", level=1, font_size=14)
toc_items = [
    ("1.", "Introduction"),
    ("2.", "Dataset Acquisition"),
    ("3.", "Data Cleaning and Pre-processing"),
    ("4.", "Exploratory Data Analysis"),
    ("5.", "Visualisations"),
    ("6.", "Critical Reflection"),
    ("7.", "References"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"  {num}  {title}")
    set_font(r, size=12)

doc.add_page_break()

# ── SECTION 1: INTRODUCTION ───────────────────────────────────────────────────

heading(doc, "1. Introduction", level=1, font_size=14)
body(doc,
     "This report presents a structured data analysis of COVID-19 statistics for the "
     "London region of the United Kingdom, covering the period from March 2020 to "
     "April 2022. The analysis was conducted using Python, employing the Pandas library "
     "for data manipulation and statistical analysis, and Matplotlib for data "
     "visualisation (McKinney, 2010; Hunter, 2007).")
body(doc,
     "The dataset was sourced from the UK Government COVID-19 Dashboard API "
     "(UK Health Security Agency, 2022), which provides publicly available, "
     "region-level epidemiological data including weekly case counts, deaths, "
     "hospital admissions, and hospital occupancy figures.")
body(doc,
     "The report is structured across five analytical phases: dataset acquisition, "
     "data cleaning and pre-processing, exploratory data analysis (including descriptive "
     "statistics, trend identification, and correlation analysis), visualisation of key "
     "insights, and a critical reflection on the analytical approach and its limitations.")

doc.add_page_break()

# ── SECTION 2: DATASET ACQUISITION ───────────────────────────────────────────

heading(doc, "2. Dataset Acquisition", level=1, font_size=14)
body(doc,
     "The dataset was retrieved programmatically using the UK Government COVID-19 "
     "Dashboard RESTful API. The API returns paginated JSON responses filtered by "
     "geographic area type and name. For this analysis, the filter parameters were set "
     "to areaType=region and areaName=London, targeting the London NHS region "
     "(area code: E12000007). The following metrics were requested: new cases by publish "
     "date, cumulative cases, new deaths within 28 days of a positive test, cumulative "
     "deaths, new hospital admissions, cumulative admissions, and current hospital cases.")
body(doc,
     "A fallback mechanism was implemented to load a representative built-in sample "
     "dataset in environments where internet connectivity is unavailable, ensuring "
     "analytical reproducibility. The final dataset contains 110 weekly records "
     "spanning March 2020 to April 2022.")

heading(doc, "2.1 Source Code — Dataset Acquisition (Unit_5_covid_analysis.py)",
        level=2, font_size=12)
code_block(doc, read_source(SRC["acquire"]))

doc.add_page_break()

# ── SECTION 3: DATA CLEANING ──────────────────────────────────────────────────

heading(doc, "3. Data Cleaning and Pre-processing", level=1, font_size=14)
body(doc,
     "Raw data ingested from public health APIs frequently contains quality issues "
     "including missing values, inconsistent data types, and unsystematic column naming "
     "conventions (Wickham, 2014). The cleaning pipeline applied to this dataset "
     "addressed each of these concerns through a sequence of deliberate pre-processing steps.")
body(doc,
     "The date column was parsed from string to datetime64 format to enable time-series "
     "operations. Column names were standardised from camelCase (e.g., "
     "newCasesByPublishDate) to snake_case (e.g., new_cases) to improve readability and "
     "consistency. Ten missing values in the deaths, admissions, and hospital columns — "
     "arising from the absence of reporting infrastructure in the earliest weeks of the "
     "pandemic — were imputed with zero, a decision justified by the contextual knowledge "
     "that tracking systems were not yet operational. Float columns were converted to "
     "integer types where decimal precision was not meaningful. Two derived analytical "
     "columns were added: case fatality rate (%) and hospitalisation rate (%), computed "
     "from cumulative totals.")

heading(doc, "3.1 Source Code — Data Cleaning (Unit_5.2_clean_data.py)",
        level=2, font_size=12)
code_block(doc, read_source(SRC["clean"]))

fig_num = 1
figure(doc, "london_covid_cleaning_overview.png",
       "Data Cleaning Overview — Missing Values Before and After Cleaning, Case and Death "
       "Trends, and Case Fatality Rate Over Time",
       fig_num)
fig_num += 1

doc.add_page_break()

# ── SECTION 4: EXPLORATORY DATA ANALYSIS ─────────────────────────────────────

heading(doc, "4. Exploratory Data Analysis", level=1, font_size=14)

heading(doc, "4.1 Descriptive Statistics", level=2, font_size=12)
body(doc,
     "Descriptive statistics were computed for five key metrics: new cases, new deaths, "
     "new admissions, hospital cases, and case fatality rate. The mean weekly case count "
     "was 24,922, while the median was substantially lower at 6,850 — a divergence "
     "explained by the pronounced positive skewness (skew = 3.94) caused by the Omicron "
     "wave peaks in December 2021 and January 2022. Similarly, the mean case fatality "
     "rate of 6.19% masks considerable temporal variation, from a high of 22.13% in the "
     "early pandemic to near-zero by 2022 following widespread vaccination.")

heading(doc, "4.2 Trend Analysis", level=2, font_size=12)
body(doc,
     "A four-week rolling average was applied to smooth short-term fluctuations and "
     "reveal underlying epidemiological trends. Three distinct transmission waves were "
     "identified: the initial wave (March–May 2020), the Alpha/Delta-driven second and "
     "third waves (October 2020–July 2021), and the Omicron surge (December 2021–February "
     "2022). Year-on-year total case counts rose from 135,245 in 2020 to 1,700,200 in "
     "2021, driven largely by improved testing capacity and the highly transmissible Delta "
     "and Omicron variants (Twohig et al., 2022). Weeks with case counts exceeding the "
     "75th percentile threshold of 29,500 were classified as high-transmission periods, "
     "totalling 28 weeks across the observation window.")

heading(doc, "4.3 Correlation Analysis", level=2, font_size=12)
body(doc,
     "Pearson correlation coefficients were computed across all metric pairs. The "
     "strongest association was observed between new deaths and hospital cases "
     "(r = 0.91), reflecting the direct clinical pathway from COVID-19 infection to "
     "hospitalisation and mortality. New admissions correlated strongly with both deaths "
     "(r = 0.77) and hospital occupancy (r = 0.76). The weaker correlation between new "
     "cases and new deaths (r = 0.27) is attributable to the decoupling effect of "
     "vaccination from 2021 onwards, which reduced infection-to-death conversion rates "
     "substantially (Haas et al., 2021). A two-week lag correlation between cases and "
     "deaths yielded r = 0.25, consistent with the known incubation and disease "
     "progression timeline.")

heading(doc, "4.4 Source Code — Full Analysis (Unit_5.3_analysis.py)",
        level=2, font_size=12)
code_block(doc, read_source(SRC["analysis"]))

heading(doc, "4.5 Source Code — Simplified Analysis (Unit_5.3.1_simplified.py)",
        level=2, font_size=12)
body(doc,
     "The following is a beginner-accessible version of the analysis, written with "
     "plain sequential logic, step-by-step comments, and no advanced Python constructs, "
     "suitable for introductory-level data science learners.")
code_block(doc, read_source(SRC["simple"]))

doc.add_page_break()

# ── SECTION 5: VISUALISATIONS ─────────────────────────────────────────────────

heading(doc, "5. Visualisations", level=1, font_size=14)
body(doc,
     "Eleven visualisations were produced using Matplotlib to communicate the analytical "
     "findings. Charts are presented below with figure captions. The source code used to "
     "generate all visualisations is provided in Section 5.1.")

body(doc,
     "Figures 2 and 3 present weekly case and death trends with rolling averages, "
     "clearly showing the three transmission waves. Figures 4 and 5 compare mean and "
     "median values, illustrating the degree to which extreme Omicron-wave weeks skew "
     "the distribution. Figures 6 and 7 examine hospital pressure over time. Figure 8 "
     "presents the Case Fatality Rate decline over time, a key indicator of the impact "
     "of the vaccination programme. Figures 9 and 10 summarise yearly totals. "
     "Figures 11 and 12 provide the Pearson correlation heatmap and pairwise bar chart. "
     "Figure 13 presents the scatter plot of deaths versus lagged cases.")

for path, caption in FIGURES:
    figure(doc, path, caption, fig_num)
    fig_num += 1

heading(doc, "5.1 Source Code — Visualisations (Unit_5.4_visualisations.py)",
        level=2, font_size=12)
code_block(doc, read_source(SRC["visuals"]))

doc.add_page_break()

# ── SECTION 6: CRITICAL REFLECTION ───────────────────────────────────────────

heading(doc, "6. Critical Reflection", level=1, font_size=14)
body(doc,
     "This analysis successfully demonstrated an end-to-end data science pipeline — from "
     "acquisition through cleaning, statistical analysis, and visualisation — applied to a "
     "real-world public health dataset. The use of Pandas enabled efficient tabular "
     "operations, while Matplotlib produced clear, publication-quality charts (McKinney, "
     "2010; Hunter, 2007). Rolling averages and correlation matrices effectively "
     "surfaced meaningful epidemiological patterns, including the three transmission waves "
     "and the progressive decoupling of case counts from mortality following vaccination.")
body(doc,
     "Nevertheless, several limitations merit acknowledgement. The dataset represents "
     "weekly aggregates at the regional level, which obscures borough-level heterogeneity "
     "within London — a city of considerable socioeconomic and demographic diversity. "
     "Imputing early missing values with zero, while contextually reasonable, "
     "introduces a minor distortion in cumulative statistics for the earliest weeks. "
     "Furthermore, the Pearson correlation assumes linearity, which may not fully capture "
     "the complex, non-linear dynamics of pandemic transmission (Haas et al., 2021).")
body(doc,
     "Future analyses should incorporate demographic stratification, borough-level "
     "granularity, and vaccination uptake rates to improve explanatory power. Machine "
     "learning models, such as ARIMA or LSTM networks, could be applied to generate "
     "short-term case forecasts, extending the utility of this dataset beyond descriptive "
     "analysis (Twohig et al., 2022). Despite these limitations, the findings provide a "
     "robust, data-driven overview of London's COVID-19 trajectory and the measurable "
     "impact of public health interventions.")

doc.add_page_break()

# ── SECTION 7: REFERENCES ─────────────────────────────────────────────────────

heading(doc, "7. References", level=1, font_size=14)

refs = [
    ("Haas, E.J., Angulo, F.J., McLaughlin, J.M., Anis, E., Singer, S.R., Khan, F., "
     "Brooks, N., Smaja, M., Mircus, G., Pan, K., Southern, J., Swerdlow, D.L., "
     "Jodar, L., Levy, Y. and Alroy-Preis, S. (2021) 'Impact and effectiveness of "
     "mRNA BNT162b2 vaccine against SARS-CoV-2 infections and COVID-19 cases, "
     "hospitalisations, and deaths following a nationwide vaccination campaign in "
     "Israel', "
     "The Lancet, 397(10287), pp. 1819–1829. "
     "Available at: https://doi.org/10.1016/S0140-6736(21)00947-8 "
     "(Accessed: 28 February 2026)."),

    ("Hunter, J.D. (2007) 'Matplotlib: A 2D graphics environment', "
     "Computing in Science and Engineering, 9(3), pp. 90–95. "
     "Available at: https://doi.org/10.1109/MCSE.2007.55 "
     "(Accessed: 28 February 2026)."),

    ("McKinney, W. (2010) 'Data structures for statistical computing in Python', "
     "Proceedings of the 9th Python in Science Conference, pp. 56–61. "
     "Available at: https://doi.org/10.25080/Majora-92bf1922-00a "
     "(Accessed: 28 February 2026)."),

    ("Sadalage, P.J. and Fowler, M. (2012) "
     "NoSQL Distilled: A Brief Guide to the Emerging World of Polyglot Persistence. "
     "Boston: Addison-Wesley."),

    ("Stonebraker, M. (2010) 'SQL databases v. NoSQL databases', "
     "Communications of the ACM, 53(4), pp. 10–11. "
     "Available at: https://doi.org/10.1145/1721654.1721659 "
     "(Accessed: 28 February 2026)."),

    ("Twohig, K.A., Nyberg, T., Zaidi, A., Thelwall, S., Sinnathamby, M.A., "
     "Aliabadi, S., Seaman, S.R., Harris, R.J., Hope, R., Lopez-Bernal, J., "
     "Charlett, A., De Angelis, D., Presanis, A.M. and Dabrera, G. (2022) "
     "'Hospital admission and emergency care attendance risk for SARS-CoV-2 delta "
     "(B.1.617.2) compared with alpha (B.1.1.7) variants of concern: a cohort study', "
     "The Lancet Infectious Diseases, 22(1), pp. 35–42. "
     "Available at: https://doi.org/10.1016/S1473-3099(21)00475-8 "
     "(Accessed: 28 February 2026)."),

    ("UK Health Security Agency (2022) Coronavirus (COVID-19) in the UK Dashboard. "
     "Available at: https://coronavirus.data.gov.uk "
     "(Accessed: 28 February 2026)."),

    ("Wickham, H. (2014) 'Tidy data', "
     "Journal of Statistical Software, 59(10), pp. 1–23. "
     "Available at: https://doi.org/10.18637/jss.v059.i10 "
     "(Accessed: 28 February 2026)."),
]

for ref in refs:
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent    = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-1.2)
    p.paragraph_format.space_after    = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    set_font(run, size=11)

# ── SAVE ──────────────────────────────────────────────────────────────────────

output = "COVID19_London_Analysis_Report.docx"
doc.save(output)
print(f"\n[SAVED] {os.path.abspath(output)}")
print(f"  Pages (approx) : will vary by Word version")
print(f"  Sections       : 7")
print(f"  Figures        : {fig_num - 1}")
print(f"  References     : {len(refs)}")
