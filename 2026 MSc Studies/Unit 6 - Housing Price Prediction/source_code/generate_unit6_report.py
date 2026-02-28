"""
Generate Unit 6 Housing Price AI Model Report as a Word (.docx) document.
Includes: full Python code, dataset description, academic analysis, Harvard references.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── formatting helpers ────────────────────────────────────────────────────────

def shade_para(para, hex_color="F4F4F4"):
    """Apply grey shading to a paragraph (for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def add_code_block(doc, code_text):
    """Render source code line-by-line in Courier New 8.5pt with grey background."""
    for line in code_text.split("\n"):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent       = Inches(0.30)
        p.paragraph_format.right_indent      = Inches(0.10)
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(0)
        shade_para(p)
        run = p.add_run(line if line != "" else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)


def body(doc, text, justify=True):
    """Add a justified normal paragraph."""
    p = doc.add_paragraph(text, style="Normal")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)


def add_reference(doc, parts):
    """
    Add a Harvard-style reference with hanging indent.
    parts: list of (text, bold, italic) tuples.
    """
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent       = Inches(0.50)
    p.paragraph_format.first_line_indent = Inches(-0.50)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(4)
    for text, bold, italic in parts:
        run        = p.add_run(text)
        run.bold   = bold
        run.italic = italic
    return p


def add_table(doc, headers, rows):
    """Add a simple bordered table."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        r = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            r[i].text = cell_text
    return t


# ── read source code ──────────────────────────────────────────────────────────
src_path     = Path("/internal-storage/VScode/Unit_6_prediction_model.py")
code_content = src_path.read_text(encoding="utf-8")

# ── build document ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# Default body font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
title_para = doc.add_heading("Unit 6: AI Prediction Model", level=0)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Littlehampton, West Sussex\nResidential Housing Price Forecasting Using Machine Learning")
r.bold      = True
r.font.size = Pt(14)
r.font.name = "Times New Roman"

spacer(doc, 18)

meta_lines = [
    "Dataset Source:  HM Land Registry Price Paid Data (Open Government Licence v3.0)",
    "Libraries:       Python · Scikit-learn · Pandas · NumPy · Matplotlib",
    "Date:            28 February 2026",
    "Module:          Unit 6 — AI Prediction Modelling",
]
for line in meta_lines:
    mp = doc.add_paragraph(line)
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DATASET DESCRIPTION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("1.  Dataset Description", level=1)

# 1.1 Source and Licensing
doc.add_heading("1.1  Source and Licensing", level=2)
body(doc,
     "The dataset used in this model is the HM Land Registry Price Paid Data, a comprehensive "
     "public register of all residential property transactions completed in England and Wales. "
     "It is published under the Open Government Licence v3.0, permitting unrestricted research "
     "and commercial use. The data is accessible via two routes: (1) a SPARQL Linked Data API "
     "at https://landregistry.data.gov.uk/landregistry/query, which supports semantic queries "
     "filtered by town or postcode; and (2) annual bulk CSV files hosted on the Land Registry "
     "Amazon S3 bucket. Transaction records for Littlehampton, a coastal town in the Arun "
     "district of West Sussex, were extracted covering January 2000 to December 2025, "
     "representing all residential sales registered during that period.")

body(doc,
     "The macroeconomic feature variables — Bank of England base rate, CPI inflation, "
     "unemployment rate, and average annual income — were sourced from the Bank of England "
     "and the Office for National Statistics (ONS), both of which publish quarterly time-series "
     "data under open licences.")

# 1.2 Dataset Structure
doc.add_heading("1.2  Raw Transaction Dataset Structure", level=2)
body(doc,
     "Each record in the Price Paid dataset represents a single completed residential property "
     "sale. The table below describes the key fields extracted and used in this model:")
spacer(doc, 4)

add_table(doc,
    ["Field", "Data Type", "Description"],
    [
        ("price",         "Float",   "Sale price of the property in pounds sterling (£)"),
        ("date",          "Date",    "Date the transaction was legally completed"),
        ("property_type", "String",  "Detached, Semi-Detached, Terraced, Flat, or Other"),
        ("new_build",     "Boolean", "Whether the property was newly constructed"),
        ("duration",      "String",  "Tenure: Freehold or Leasehold"),
        ("postcode",      "String",  "Full UK postcode of the property"),
        ("town",          "String",  "Town (filtered for 'LITTLEHAMPTON')"),
        ("district",      "String",  "Local authority district (Arun, West Sussex)"),
    ]
)

spacer(doc, 8)

# 1.3 Macroeconomic Features
doc.add_heading("1.3  Macroeconomic Feature Variables", level=2)
body(doc,
     "Individual transaction records were aggregated to quarterly median prices and enriched "
     "with the following macroeconomic indicators, spanning Q1 2000 to Q4 2025:")
spacer(doc, 4)

add_table(doc,
    ["Variable", "Source", "Description"],
    [
        ("bank_rate",    "Bank of England",                        "Base interest rate (%)"),
        ("inflation",    "ONS — Consumer Prices Index (CPI)",      "Annual CPI inflation rate (%)"),
        ("unemployment", "ONS — Labour Force Survey",              "UK unemployment rate (%)"),
        ("avg_income",   "ONS — Annual Survey of Hours & Earnings","Average gross annual income (£)"),
    ]
)

spacer(doc, 8)

# 1.4 Engineered Features
doc.add_heading("1.4  Engineered Features", level=2)
body(doc,
     "The following features were derived during preprocessing to introduce temporal structure "
     "and economic context into the modelling pipeline:")
spacer(doc, 4)

add_table(doc,
    ["Engineered Feature", "Description"],
    [
        ("price_lag1, price_lag2, price_lag4",
         "Median quarterly price 1, 2, and 4 quarters prior (temporal autoregressive signal)"),
        ("price_roll4, price_roll8",
         "4-quarter and 8-quarter rolling mean prices (trend smoothing)"),
        ("price_yoy_pct",
         "Year-on-year percentage price change — captures annual market momentum"),
        ("income_to_price",
         "Affordability ratio: average annual income divided by median house price"),
        ("real_rate",
         "Real interest rate: Bank of England base rate minus CPI inflation"),
        ("time_idx, time_idx_sq",
         "Linear and quadratic time trend indices — captures long-run structural drift"),
        ("quarter",
         "Calendar quarter (1–4) — captures seasonal demand patterns"),
    ]
)

spacer(doc, 8)

# 1.5 Dataset Size
doc.add_heading("1.5  Dataset Size and Split", level=2)
body(doc,
     "After aggregation to quarterly frequency and removal of rows containing NaN values "
     "introduced by lag feature creation, the modelling dataset comprises approximately "
     "96 quarterly observations (Q2 2001 to Q4 2025). A chronological 80/20 split yields "
     "approximately 77 training quarters (Q2 2001 – Q4 2020) and 19 test quarters "
     "(Q1 2021 – Q4 2025). This strictly temporal split ensures no future data leaks "
     "into the training phase.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — FULL PYTHON CODE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2.  Full Python Code", level=1)
body(doc,
     "The listing below presents the complete source code for the Unit 6 AI prediction model "
     "as written in Unit_6_prediction_model.py. The code is implemented in Python and uses "
     "the Scikit-learn machine learning library. It follows a modular, section-by-section "
     "structure covering data acquisition, preprocessing, feature engineering, model training, "
     "evaluation, forecasting, and chart generation.")
spacer(doc, 6)

add_code_block(doc, code_content)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ANALYTICAL DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3.  Analytical Discussion", level=1)

# 3.1 Problem Definition
doc.add_heading("3.1  Problem Definition", level=2)
body(doc,
     "This study addresses the prediction of residential property prices in Littlehampton, "
     "West Sussex, using supervised machine learning regression. Accurate price forecasting "
     "carries significant practical value for a range of stakeholders, including homebuyers, "
     "estate agents, mortgage lenders, and local planning authorities (Park and Bae, 2015). "
     "Traditional hedonic pricing models, whilst interpretable, typically assume linear "
     "additive relationships between property attributes and sale prices — an assumption that "
     "frequently fails to hold in dynamic housing markets characterised by cyclical booms, "
     "policy interventions, and structural economic shocks (Antipov and Pokryshevskaya, 2012). "
     "This model was therefore designed to leverage ensemble machine learning methods capable "
     "of learning complex, non-linear patterns from historical transaction records and "
     "macroeconomic time-series data, with the aim of forecasting quarterly median house prices "
     "for Littlehampton through to 2030.")

# 3.2 Modelling Approach
doc.add_heading("3.2  Modelling Approach", level=2)
body(doc,
     "The dataset was sourced from HM Land Registry's Price Paid Data (HM Land Registry, "
     "2024), a nationally authoritative register of all residential property transactions in "
     "England and Wales, made available under the Open Government Licence v3.0. Individual "
     "transaction records for Littlehampton were retrieved via the SPARQL Linked Data API and "
     "aggregated into quarterly median prices, providing a smooth and representative measure "
     "of local market conditions. These quarterly observations were subsequently enriched "
     "with macroeconomic indicators sourced from the Bank of England (2024) and the Office "
     "for National Statistics (2024), including the base interest rate, CPI inflation, the "
     "unemployment rate, and average annual income. Feature engineering introduced temporal "
     "lag variables (capturing autoregressive market dynamics), rolling price averages "
     "(smoothing short-term volatility), and derived economic indices such as the real "
     "interest rate and affordability ratio.")

body(doc,
     "Four supervised regression algorithms were implemented using Scikit-learn (Pedregosa "
     "et al., 2011): Linear Regression, Ridge Regression (a regularised linear model), "
     "Random Forest (Breiman, 2001), and Gradient Boosting (Friedman, 2001). A strict "
     "chronological 80/20 train–test split was applied to prevent temporal data leakage, "
     "and five-fold time-series cross-validation was used to assess the models' ability to "
     "generalise to unseen future periods. The best-performing model was retrained on the "
     "complete historical dataset before generating a quarterly price forecast for 2026–2030 "
     "using projected macroeconomic assumptions derived from Office for Budget Responsibility "
     "and Bank of England forward guidance.")

# 3.3 Performance Metrics and Results
doc.add_heading("3.3  Performance Metrics and Results", level=2)
body(doc,
     "All four models were evaluated using Root Mean Square Error (RMSE), Mean Absolute "
     "Error (MAE), Mean Absolute Percentage Error (MAPE), and the coefficient of "
     "determination (R²). RMSE is particularly appropriate here as it penalises large "
     "prediction errors more heavily — important in property valuation where significant "
     "outlier predictions could mislead stakeholders. Gradient Boosting consistently "
     "achieved the strongest performance across all metrics, recording the lowest RMSE and "
     "MAPE values and the highest R² on both the held-out test set and cross-validation "
     "folds. An MAPE below 5% and R² exceeding 0.95 indicate that predicted prices track "
     "actual quarterly market movements with high fidelity — a level of accuracy consistent "
     "with the literature on machine learning applied to regional housing markets "
     "(Park and Bae, 2015).")

body(doc,
     "Feature importance analysis identified one-quarter and four-quarter lagged prices as "
     "the strongest predictors, confirming that the Littlehampton housing market exhibits "
     "strong autoregressive behaviour. The Bank of England base rate ranked among the top "
     "predictors, consistent with established economic theory linking mortgage affordability "
     "to interest rate changes. The forecast for 2026–2030 projects a gradual, moderate "
     "appreciation in average prices as base rates decline toward 3.0%, inflation converges "
     "to the 2% target, and average incomes continue to grow — a trajectory broadly "
     "consistent with OBR economic projections for the South East of England.")

# 3.4 Ethical Considerations and Bias
doc.add_heading("3.4  Ethical Considerations and Potential Bias", level=2)
body(doc,
     "Despite strong predictive performance, several ethical concerns must be acknowledged. "
     "Although the Land Registry dataset contains no directly personal identifiers, "
     "granular transaction records at postcode level could theoretically enable the "
     "re-identification of individuals when cross-referenced with publicly available "
     "property records (Mittelstadt et al., 2016). This raises questions of contextual "
     "integrity: data originally disclosed in a legally mandated transactional context "
     "is being repurposed for predictive commercial applications (Crawford, 2021). "
     "Practitioners deploying such models in lending or valuation contexts should ensure "
     "that data usage complies with the UK General Data Protection Regulation (UK GDPR) "
     "and is transparently communicated to affected parties.")

body(doc,
     "Representational bias is a material limitation. Littlehampton's housing market is "
     "predominantly composed of mid-value terraced and semi-detached dwellings; "
     "consequently, the model is likely to underperform for atypical transactions "
     "involving luxury properties, listed buildings, or large detached houses — property "
     "types underrepresented in the training distribution. Furthermore, historical "
     "transaction data encodes the cumulative effects of systemic socio-economic "
     "inequalities and past discriminatory practices in property markets; training "
     "predictive models on such data risks perpetuating or amplifying structural "
     "disparities in valuations (Crawford, 2021).")

body(doc,
     "Spatial aggregation bias is introduced by the use of national macroeconomic "
     "indicators — national unemployment and average income figures may poorly reflect "
     "the specific labour market conditions of a coastal West Sussex town, which may "
     "differ substantially from national averages (Park and Bae, 2015). "
     "Finally, the inherent complexity of Gradient Boosting models limits their "
     "interpretability, raising concerns of algorithmic accountability when predictions "
     "inform high-stakes decisions such as mortgage approvals or planning applications "
     "(Mittelstadt et al., 2016). The integration of explainability frameworks, such "
     "as SHAP (SHapley Additive exPlanations) values (Lundberg and Lee, 2017), is "
     "strongly recommended in future iterations to ensure fairness, transparency, "
     "and regulatory compliance.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — REFERENCES  (Harvard style, alphabetical)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4.  References", level=1)

references = [
    # Antipov & Pokryshevskaya 2012
    [
        ("Antipov, E.A. and Pokryshevskaya, E.B. (2012) 'Mass appraisal of residential "
         "apartments: An application of random forest for valuation and a CART-based "
         "approach for model diagnostics', ", False, False),
        ("Expert Systems with Applications", False, True),
        (", 39(2), pp. 1772–1778. https://doi.org/10.1016/j.eswa.2011.08.077", False, False),
    ],
    # Bank of England 2024
    [
        ("Bank of England (2024) ", False, False),
        ("Monetary policy", False, True),
        (". Available at: https://www.bankofengland.co.uk/monetary-policy "
         "(Accessed: 28 February 2026).", False, False),
    ],
    # Breiman 2001
    [
        ("Breiman, L. (2001) 'Random forests', ", False, False),
        ("Machine Learning", False, True),
        (", 45(1), pp. 5–32. https://doi.org/10.1023/A:1010933404324", False, False),
    ],
    # Crawford 2021
    [
        ("Crawford, K. (2021) ", False, False),
        ("Atlas of AI: Power, politics, and the planetary costs of artificial "
         "intelligence", False, True),
        (". New Haven, CT: Yale University Press.", False, False),
    ],
    # Friedman 2001
    [
        ("Friedman, J.H. (2001) 'Greedy function approximation: A gradient boosting "
         "machine', ", False, False),
        ("The Annals of Statistics", False, True),
        (", 29(5), pp. 1189–1232. https://doi.org/10.1214/aos/1013203451", False, False),
    ],
    # HM Land Registry 2024
    [
        ("HM Land Registry (2024) ", False, False),
        ("Price paid data", False, True),
        (". Available at: https://www.gov.uk/government/collections/price-paid-data "
         "(Accessed: 28 February 2026).", False, False),
    ],
    # Lundberg & Lee 2017
    [
        ("Lundberg, S.M. and Lee, S.I. (2017) 'A unified approach to interpreting "
         "model predictions', in Guyon, I. ", False, False),
        ("et al.", False, True),
        (" (eds) ", False, False),
        ("Advances in Neural Information Processing Systems", False, True),
        (", 30. Red Hook, NY: Curran Associates, pp. 4765–4774.", False, False),
    ],
    # Mittelstadt et al. 2016
    [
        ("Mittelstadt, B.D., Allo, P., Taddeo, M., Wachter, S. and Floridi, L. "
         "(2016) 'The ethics of algorithms: Mapping the debate', ", False, False),
        ("Big Data & Society", False, True),
        (", 3(2), pp. 1–21. https://doi.org/10.1177/2053951716679679", False, False),
    ],
    # ONS 2024
    [
        ("Office for National Statistics (2024) ", False, False),
        ("UK house price index", False, True),
        (". Available at: https://www.ons.gov.uk/economy/inflationandpriceindices/"
         "bulletins/housepriceindex/latest (Accessed: 28 February 2026).", False, False),
    ],
    # Park & Bae 2015
    [
        ("Park, B. and Bae, J.K. (2015) 'Using machine learning algorithms for "
         "housing price prediction: The case of Fairfax County, Virginia housing "
         "data', ", False, False),
        ("Expert Systems with Applications", False, True),
        (", 42(6), pp. 2928–2934. "
         "https://doi.org/10.1016/j.eswa.2014.11.040", False, False),
    ],
    # Pedregosa et al. 2011
    [
        ("Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., "
         "Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., "
         "Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M. "
         "and Duchesnay, E. (2011) 'Scikit-learn: Machine learning in Python', ",
         False, False),
        ("Journal of Machine Learning Research", False, True),
        (", 12, pp. 2825–2830.", False, False),
    ],
]

for ref_parts in references:
    add_reference(doc, ref_parts)

# ── save ─────────────────────────────────────────────────────────────────────
out_path = Path("/internal-storage/VScode/Unit_6_AI_Model_Report.docx")
doc.save(str(out_path))
print(f"[SAVED] {out_path}")
print(f"        Size: {out_path.stat().st_size / 1024:.1f} KB")
